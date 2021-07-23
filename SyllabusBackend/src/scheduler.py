import docx
from datetime import datetime
from .colors import Color
from json import load
from typing import List, Dict, Union

# for testing
from inspect import currentframe
from os.path import dirname, realpath

MAIN_FONT: str = "Times New Roman"

MIN_COLS: int = 2
MAX_COLS: int = 6

# if the reworked holiday or other important deadline something falls on a date in the system, it goes in Due.
# if not, it goes on the day before and states when the deadline is

class MeetingSchedule:
    def __init__(self, document: docx.Document = None):
        self._document: docx.Document = document

    def build(self, json_filename: str, return_schedule: bool = True) -> str:
        print("%-25s %-30s %10s" % (__name__, currentframe().f_code.co_name, "Building meeting table..."))

        # retrieve json dict from frontend
        dict_from_json = None
        with open(json_filename) as file:
            dict_from_json = load(file)

        # if no document was passed then we're dealing with a standalone schedule
        # rather than adding one into a syllabus
        if not self._document:
            self._document = docx.Document()
            self._style_set(self._document)

        # for now im only allowing landscape functionality if the schedule is a different document
        if dict_from_json["landscape"] and not return_schedule: self._flip_orientation()

        if not return_schedule:
            self._draw_standlone_requirements(
                course=dict_from_json["courseName"],
                sem=dict_from_json["semester"],
                instructor=dict_from_json["instructorName"]
            )

        self._gen_meet_table(content=dict_from_json["content"], num_cols=dict_from_json["numCols"])

        print("%-25s %-30s %5s" % (__name__, currentframe().f_code.co_name, f"{Color.OKCYAN}Done!{Color.ENDC}"))
        if return_schedule: return self._document

        self._document.save(dict_from_json["filename"])
        print("%10s %62s" % (f"{Color.OKCYAN}Path to schedule:", f"%s/%s{Color.ENDC}" % (dirname(realpath(dict_from_json["filename"])), dict_from_json["filename"][5:])))

        return

    def _draw_standlone_requirements(self, course: str, sem: str, instructor: str) -> None:
        # generate course name and semester heading
        heading_text: str = "%s - %s\n%s" % (course, sem, instructor)
        heading: docx.text.paragraph.Paragraph = self._document.add_paragraph(heading_text, style="Header")
        heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        # draw the last updated header
        curr_time: str = "%s/%s/%s" % (datetime.now().month, datetime.now().day, datetime.now().year)
        last_updated_header: docx.text.paragraph.Paragraph = self._document.sections[0].header.paragraphs[0]
        last_updated_header.text = "Last updated: %s" % curr_time
        last_updated_header.style = self._document.styles["SYL_TEXT"]
        last_updated_header.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

        return

    def _flip_orientation(self) -> None:
        print("%-25s %-30s %10s" % (__name__, currentframe().f_code.co_name, "Flipping orientation..."))

        curr_section: docx.section.Section = self._document.sections[0]
        curr_section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE

        old_height: docx.shared.Twips = curr_section.page_height
        curr_section.page_height = curr_section.page_width
        curr_section.page_width = old_height

        return

    # RIPPED THIS STRAIGHT FROM SYLLABUS FOR USE HERE
    # create new styles that parts of the syllabus will fall under
    def _style_set(self, document) -> None:
        print("%-25s %-30s %10s" % (__name__, currentframe().f_code.co_name, "Setting schedule document styles..."))

        # Set font and size of the very first heading of the syllabus
        document.styles["Header"].font.name = MAIN_FONT
        document.styles["Header"].font.size = docx.shared.Pt(25)

        # Set font and the size for the header of every section
        headers: docx.styles.style._ParagraphStyle = document.styles.add_style("SYL_HEAD", docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        headers.base_style = document.styles["Heading 1"]
        headers.font.name = MAIN_FONT
        headers.font.size = docx.shared.Pt(16)

        # Set font and size for the text under each header
        runs: docx.styles.style._ParagraphStyle = document.styles.add_style("SYL_TEXT", docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        runs.base_style = document.styles["Normal"]
        runs.font.name = MAIN_FONT
        runs.font.size = docx.shared.Pt(12)

        # Set font, size, and table aesthetics
        table: docx.styles.style._TableStyle = document.styles.add_style("SYL_TABLE", docx.enum.style.WD_STYLE_TYPE.TABLE)
        table.base_style = document.styles["Medium Grid 1"]
        table.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        table.paragraph_format.line_spacing = 1
        table.font.name = MAIN_FONT
        table.font.size = docx.shared.Pt(10)

        return

    # priv helper for setting up a new section of the syllabus via a header
    def _gen_header(self, header_title: str) -> docx.text.paragraph.Paragraph:
        heading: docx.text.paragraph.Paragraph = self._document.add_paragraph(header_title, style="SYL_HEAD")
        heading.paragraph_format.line_spacing = 1

        return heading

    def _gen_meet_table(self, content: Dict[str, List[str]], num_cols: int = 3):
        print("%-25s %-30s %10s" % (__name__, currentframe().f_code.co_name, "Generating the meeting table..."))

        if not MIN_COLS <= num_cols <= MAX_COLS: raise ValueError(f"{Color.FAIL}Column Error.{Color.ENDC}")
        num_rows: int = len(content["Date"])

        # move Registrar info and Religious Observances to the end of the dictionary
        # so they appear as the rightmost columns in the table
        self._reorder_content(content, "Registrar Info")
        self._reorder_content(content, "Common Religious Observances")

        # generate course scheduling title
        self._gen_header("Course Schedule")

        table: docx.table.Table = self._document.add_table(rows=num_rows + 1, cols=num_cols, style="Table Grid")

        print("%-25s %-30s %10s" % (__name__, currentframe().f_code.co_name, "Populating table with fields from JSON..."))

        # i: column index
        # key: the column label
        # value: what the user wants entered in the table
        for i, (key, value) in enumerate(content.items()):
            print(i, key)
            table.rows[0].cells[i].text = key
            # make me a header vvv
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

            for j in range(len(value)):
                if key == "Date":
                    table.rows[j + 1].cells[0].text = value[j]
                    table.rows[j + 1].cells[i].paragraphs[0].runs[0].bold = True

                if key == "Registrar Info":
                    table.rows[j + 1].cells[i].text = value[j]

                if key == "Common Religious Observances":
                    table.rows[j + 1].cells[i].text = value[j]

        table.autofit = True

        return

    # this just takes a desired key in a dict and puts it at the end of it
    def _reorder_content(self, content: Dict[str, List[str]], key: str) -> None:
        content_keys: List[str] = list(content.keys())

        if key in content_keys:
            info: List[str] = content[key]
            del content[key]

            content[key] = info
        return
