import docx
from bs4 import BeautifulSoup, element
from .colors import Color
from datetime import datetime
from docxcompose.composer import Composer
from os import system
from .dates import Dates
from pypandoc import convert_text
from requests import get as req_get
from .scheduler import MeetingSchedule
from typing import List

# for testing
from inspect import currentframe
from os.path import dirname, realpath

MAIN_FONT: str = "Times New Roman"

class DOCXWriter:
    def __init__(self, dates_obj: Dates):
        self._document: docx.Document = docx.Document()
        self._dates_obj: Dates = dates_obj
        self._style_set(self._document)

    def populate_title(self, course_name: str) -> None:
        print("%-25s %-30s %10s" % (__name__, currentframe().f_code.co_name, "Populating syllabus header..."))

        # put some meeting times under the header maybe?
        # prob could go under course info too?
        heading: docx.text.paragraph.Paragraph = self._document.add_paragraph(course_name, style="Header")
        #heading.style.font.size = docx.shared.Pt(25)
        heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        return

    def populate_course_info(self, name: str, email: str, oh_info: str) -> None:
        print("%-25s %-30s %10s" % (__name__, currentframe().f_code.co_name, "Populating course info..."))

        self._gen_header("Course Information")
        para: docx.text.paragraph.Paragraph = self._gen_paragraph("")

        para.add_run("Instructor: ").bold = True
        para.add_run("%s\n" % name)

        para.add_run("E-Mail: ").bold = True
        para.add_run("%s\n" % email)

        para.add_run("Office Hours: ").bold = True
        para.add_run(oh_info)

        return

    def populate_assistant_info(self, name: str, email: str, oh_info: str, tf: bool = True) -> None:
        print("%-25s %-30s %10s" % (__name__, currentframe().f_code.co_name, "Populating assistant info..."))

        position: str = "TF" if tf else "TA"
        para: docx.text.paragraph.Paragraph = self._gen_paragraph("")

        para.add_run("%s: " % position).bold = True
        para.add_run("%s\n" % name)

        para.add_run("E-Mail: ").bold = True
        para.add_run("%s\n" % email)

        para.add_run("Office Hours: ").bold = True
        para.add_run(oh_info)

        return

    def populate_generic_field(self, header: str, body: str) -> None:
        print("%-25s %-30s %10s" % (__name__, currentframe().f_code.co_name, "Populating \"%s\"...") % header)

        self._gen_header(header)
        self._gen_paragraph(body)

        return

    def populate_grade_dist(self, fields: dict, raw_vals: bool = True) -> None:
        print("%-25s %-30s %10s" % (__name__, currentframe().f_code.co_name, "Populating grade distribution..."))

        self._gen_header("Grade Distribution")

        table: docx.table.Table = self._document.add_table(rows=len(fields.keys()) + 1, cols=2, style="SYL_TABLE")

        table.rows[0].cells[0].text = "Category"
        table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[0].cells[1].text = "Points" if raw_vals else "Percent"
        table.rows[0].cells[1].paragraphs[0].runs[0].bold = True

        for i, (key, val) in enumerate(fields.items()):
            table.rows[i + 1].cells[0].text = key
            table.rows[i + 1].cells[1].text = val if raw_vals else str(val) + '%'

        return

    def populate_grade_scale(self, max_points: int) -> None:
        print("%-25s %-30s %10s" % (__name__, currentframe().f_code.co_name, "Populating grade scale..."))

        self._gen_header("Grade Scale")

        table: docx.table.Table = self._document.add_table(rows=6, cols=2, style="SYL_TABLE")

        table.rows[0].cells[0].text = "Points Earned"
        table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[0].cells[1].text = "Grade"
        table.rows[0].cells[1].paragraphs[0].runs[0].bold = True

        # there is definitely a smarter way to do this
        grade_a: str = "%d - %d" % (max_points * .9, max_points)
        grade_b: str = "%d - %.2f" % (max_points * .8, max_points * .9 - .01)
        grade_c: str = "%d - %.2f" % (max_points * .7, max_points * .8 - .01)
        grade_d: str = "%d - %.2f" % (max_points * .6, max_points * .7 - .01)
        grade_f: str = "<= %.2f" % (max_points * .6 - .01)

        table.rows[1].cells[0].text = grade_a
        table.rows[1].cells[1].text = "A"
        table.rows[2].cells[0].text = grade_b
        table.rows[2].cells[1].text = "B"
        table.rows[3].cells[0].text = grade_c
        table.rows[3].cells[1].text = "C"
        table.rows[4].cells[0].text = grade_d
        table.rows[4].cells[1].text = "D"
        table.rows[5].cells[0].text = grade_f
        table.rows[5].cells[1].text = "F"

        return

    def write(self, filename: str, json_filename: str = "", separate_doc: bool = False) -> None:
        print("%-25s %-30s %10s" % (__name__, currentframe().f_code.co_name, "Saving syllabus to .docx..."))

        self._populate_meetings(json_filename, separate_doc)

        self._last_updated()

        self._populate_req_lang(filename)
        self._format_oei_lang()
        self._document.save(filename)

        print("%-25s %-30s %10s" % (__name__, currentframe().f_code.co_name, "Deleting __I_CONTAIN_UNFORMATTED_OEI_LANGUAGE__.docx..."))
        system("rm __I_CONTAIN_UNFORMATTED_OEI_LANGUAGE__.docx")

        print("%-25s %-30s %5s" % (__name__, currentframe().f_code.co_name, f"{Color.OKCYAN}Done!{Color.ENDC}"))
        print("%10s %62s" % (f"{Color.OKCYAN}Path to syllabus:", f"%s/%s{Color.ENDC}" % (dirname(realpath(filename)), filename[12:])))

        return

    def _populate_meetings(self, json_filename: str, separate_doc: bool) -> None:
        print("%-25s %-30s %10s" % (__name__, currentframe().f_code.co_name, "Populating meeting table..."))

        if separate_doc:
            self._document = MeetingSchedule(self._document).build(json_filename, separate_doc)
        else:
            MeetingSchedule().build(json_filename, separate_doc)

        return

    def _last_updated(self) -> None:
        print("%-25s %-30s %10s" % (__name__, currentframe().f_code.co_name, "Drawing last updated header..."))

        curr_time: str = "%s/%s/%s" % (datetime.now().month, datetime.now().day, datetime.now().year)
        last_updated_header: docx.text.paragraph.Paragraph = self._document.sections[0].header.paragraphs[0]
        last_updated_header.text = "Last updated: %s" % curr_time
        last_updated_header.style = self._document.styles["SYL_TEXT"]
        last_updated_header.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

    def _populate_req_lang(self, filename: str) -> None:
        print("%-25s %-30s %10s" % (__name__, currentframe().f_code.co_name, "Populating required language..."))

        oei_doc: str = "__I_CONTAIN_UNFORMATTED_OEI_LANGUAGE__.docx"
        convert_text(self._fetch_oei_lang(), "docx", format="html", outputfile=oei_doc)
        oei_doc: docx.Document = docx.Document(oei_doc)

        syllabus: Composer = Composer(self._document)
        syllabus.append(oei_doc)
        syllabus.save(filename)

        return

    def _fetch_oei_lang(self) -> element.Tag:
        print("%-25s %-30s %10s" % (__name__, currentframe().f_code.co_name, "Fetching OEI language..."))

        page = req_get("https://oei.umbc.edu/sample-title-ix-responsible-employee-syllabus-language/")
        if page.status_code != 200: raise AttributeError(f"{Color.FAIL}Could not find OEI language.")

        soup = BeautifulSoup(page.content, "html.parser")
        article = soup.find("article")

        print("%-25s %-30s %5s" % (__name__, currentframe().f_code.co_name, f"{Color.OKCYAN}Done!{Color.ENDC}"))

        return article

    def _format_oei_lang(self) -> None:
        print("%-25s %-30s %10s" % (__name__, currentframe().f_code.co_name, "Formatting OEI language..."))

        for para in self._document.paragraphs:
            if para.style.name not in ["SYL_HEAD", "SYL_TEXT", "Header"]:
                #if para.style.name
                if para.style.name == "Heading 4":
                    para.style = self._document.styles["SYL_HEAD"]
                    para.paragraph_format.line_spacing = 1

                    if "(required)" in para.text:
                        para.text = para.text.split(" (required)")[0]

                elif para.style.name in ["First Paragraph", "Body Text", "Compact"]:
                    para.style = self._document.styles["SYL_TEXT"]
                    para.style.font.name = MAIN_FONT
                    para.style.font.size = docx.shared.Pt(12)

                else:
                    para_elem = para._element
                    para_elem.getparent().remove(para_elem)
                    para_elem._p = para_elem._element = None

    # create new styles that parts of the syllabus will fall under
    def _style_set(self, document) -> None:
        print("%-25s %-30s %10s" % (__name__, currentframe().f_code.co_name, "Setting document styles..."))

        # Set font and size of the very first heading of the syllabus
        document.styles["Header"].font.name = MAIN_FONT
        document.styles["Header"].font.size = docx.shared.Pt(25)

        # Set font and the size for the header of every section
        headers: docx.styles.style._ParagraphStyle = document.styles.add_style("SYL_HEAD", docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        headers.base_style = document.styles["Heading 1"]
        headers.font.name = MAIN_FONT
        headers.font.size = docx.shared.Pt(16)

        # Set font and size for the text under each header
        runs: docx.styles.style._ParagraphStyle = document.styles.add_style("SYL_TEXT", docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        runs.base_style = document.styles["Normal"]
        runs.font.name = MAIN_FONT
        runs.font.size = docx.shared.Pt(12)

        # Set font, size, and table aesthetics
        table: docx.styles.style._TableStyle = document.styles.add_style("SYL_TABLE", docx.enum.style.WD_STYLE_TYPE.TABLE)
        table.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        table.paragraph_format.line_spacing = 1
        table.font.name = MAIN_FONT
        table.font.size = docx.shared.Pt(10)

    # priv helper for setting up a new section of the syllabus via a header
    def _gen_header(self, header_title: str) -> docx.text.paragraph.Paragraph:
        heading: docx.text.paragraph.Paragraph = self._document.add_paragraph(header_title, style="SYL_HEAD")
        heading.paragraph_format.line_spacing = 1

        return heading

    # priv helper for populating text under a header
    def _gen_paragraph(self, paragraph_text: str) -> docx.text.paragraph.Paragraph:
        paragraph: docx.text.paragraph.Paragraph = self._document.add_paragraph(paragraph_text, style="SYL_TEXT")
        paragraph.paragraph_format.line_spacing = 1

        return paragraph
